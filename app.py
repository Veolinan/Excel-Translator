import streamlit as st
import pandas as pd
from deep_translator import GoogleTranslator
import io
import time
import base64
import zipfile
from docx import Document
import fitz  # PyMuPDF
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from streamlit_local_storage import LocalStorage

# --- 1. LANGUAGE MAPPING ---
LANG_MAP = {
    "Auto Detect": "auto", "English": "en", "French": "fr", "Spanish": "es", 
    "German": "de", "Italian": "it", "Portuguese": "pt", "Chinese (Simplified)": "zh-CN",
    "Japanese": "ja", "Korean": "ko", "Russian": "ru", "Arabic": "ar", 
    "Hindi": "hi", "Turkish": "tr", "Dutch": "nl", "Greek": "el"
}

# --- 2. CONFIGURATION & BROWSER STORAGE ---
st.set_page_config(page_title="Translet Pro", layout="wide", page_icon="🌍")

# Initialize LocalStorage with correct methods
local_storage = LocalStorage()

if "step" not in st.session_state: st.session_state.step = "welcome"
if "history" not in st.session_state: st.session_state.history = []

# Load cache from Browser LocalStorage using getItem
if "local_cache" not in st.session_state:
    try:
        # The library uses getItem for retrieval
        cached_data = local_storage.getItem("translet_cache")
        st.session_state.local_cache = cached_data if cached_data else {}
    except:
        st.session_state.local_cache = {}

# Premium CSS
st.markdown("""
    <style>
    .main { background-color: #f8f9fa; }
    div.stButton > button:first-child {
        background-color: #007AFF; color: white; border-radius: 12px;
        padding: 0.6rem 2rem; font-weight: 600; border: none;
    }
    .splash-card {
        background: white; padding: 3rem; border-radius: 25px;
        box-shadow: 0 15px 35px rgba(0,0,0,0.05); text-align: center;
        max-width: 800px; margin: auto;
    }
    </style>
    """, unsafe_allow_html=True)

# --- 3. HELPERS ---

def sync_cache_to_browser():
    """Saves translation cache to browser using setItem."""
    try:
        local_storage.setItem("translet_cache", st.session_state.local_cache)
    except:
        pass

def trigger_auto_download(filename, data, mime):
    b64 = base64.b64encode(data).decode()
    dl_link = f"<script>var a=document.createElement('a');a.href='data:{mime};base64,{b64}';a.download='{filename}';document.body.appendChild(a);a.click();document.body.removeChild(a);</script>"
    st.components.v1.html(dl_link, height=0)

def translate_block(text, src, target):
    if not text or len(str(text).strip()) < 2: return text
    clean_text = str(text).strip()
    cache_key = f"{src}-{target}:{clean_text}"
    
    if cache_key in st.session_state.local_cache:
        return st.session_state.local_cache[cache_key]
    
    try:
        time.sleep(0.05) 
        res = GoogleTranslator(source=src, target=target).translate(clean_text)
        st.session_state.local_cache[cache_key] = res
        sync_cache_to_browser() 
        return res
    except:
        return text

# --- 4. SIDEBAR ---
with st.sidebar:
    st.title("🌍 Translet")
    st.markdown("---")
    st.subheader("Navigation")
    if st.button("🏠 Home", use_container_width=True):
        st.session_state.step = "welcome"; st.rerun()
    if st.button("📂 New Translation", use_container_width=True):
        st.session_state.step = "upload"; st.rerun()
    
    st.markdown("---")
    st.subheader("📜 Session History")
    
    if st.session_state.history:
        # ZIP All feature
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
            for item in st.session_state.history:
                zip_file.writestr(f"translet_{item['name']}", item['data'])
        
        st.download_button("📦 Download All as ZIP", data=zip_buffer.getvalue(), file_name="translet_batch.zip", use_container_width=True)
        
        for item in reversed(st.session_state.history):
            with st.expander(f"📄 {item['name'][:15]}..."):
                st.caption(f"To: {item['target_name']}")
                st.download_button("Download", item['data'], f"translet_{item['name']}", item['mime'], key=f"h_{item['timestamp']}")
    else:
        st.caption("No history yet.")
    
    if st.session_state.history and st.button("🗑️ Clear History", use_container_width=True):
        st.session_state.history = []; st.rerun()

# --- 5. MULTI-SCREEN UI ---

if st.session_state.step == "welcome":
    st.markdown('<div style="height:10vh;"></div>', unsafe_allow_html=True)
    st.markdown("""<div class="splash-card">
        <h1 style='font-size: 4rem; background: -webkit-linear-gradient(#007AFF, #00C7BE); -webkit-background-clip: text; -webkit-text-fill-color: transparent;'>Translet</h1>
        <p style='font-size: 1.2rem; color: #6e6e73;'>Professional document translation with persistent browser caching.</p>
    </div>""", unsafe_allow_html=True)
    _, col2, _ = st.columns([1, 0.6, 1])
    with col2:
        if st.button("Start Now", use_container_width=True):
            st.session_state.step = "upload"; st.rerun()

elif st.session_state.step == "upload":
    st.title("📂 New Translation Task")
    col_left, col_right = st.columns([1.2, 1], gap="large")
    with col_left:
        st.subheader("1. Source Documents")
        files = st.file_uploader("Upload Files", type=["xlsx", "csv", "docx", "pdf"], accept_multiple_files=True)
    with col_right:
        st.subheader("2. Language Settings")
        src_label = st.selectbox("Source Language", options=list(LANG_MAP.keys()), index=0)
        target_label = st.selectbox("Target Language", options=[k for k in LANG_MAP.keys() if k != "Auto Detect"], index=0)
        if files:
            if st.button("🚀 Process Batch", use_container_width=True):
                st.session_state.files = files
                st.session_state.src_code = LANG_MAP[src_label]
                st.session_state.target_code = LANG_MAP[target_label]
                st.session_state.target_name = target_label
                st.session_state.step = "processing"; st.rerun()

elif st.session_state.step == "processing":
    
    st.title("⚙️ Translating...")
    results = {}
    progress_bar = st.progress(0)
    
    for idx, file in enumerate(st.session_state.files):
        with st.status(f"Translating `{file.name}`...", expanded=True) as status:
            if file.name.endswith(".docx"):
                doc = Document(file)
                for p in doc.paragraphs:
                    if p.text.strip(): p.text = translate_block(p.text, st.session_state.src_code, st.session_state.target_code)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if p.text.strip(): p.text = translate_block(p.text, st.session_state.src_code, st.session_state.target_code)
                out = io.BytesIO(); doc.save(out)
                data, mime = out.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

            elif file.name.endswith(".pdf"):
                pdf_doc = fitz.open(stream=file.read(), filetype="pdf")
                packet = io.BytesIO(); can = canvas.Canvas(packet, pagesize=letter)
                for page in pdf_doc:
                    text = page.get_text()
                    trans = translate_block(text, st.session_state.src_code, st.session_state.target_code)
                    t = can.beginText(50, 750); t.setFont("Helvetica", 10)
                    for line in trans.split('\n'): t.textLine(line[:100])
                    can.drawText(t); can.showPage()
                can.save()
                data, mime = packet.getvalue(), "application/pdf"

            elif file.name.endswith((".csv", ".xlsx")):
                df = pd.read_csv(file) if file.name.endswith(".csv") else pd.read_excel(file)
                cols = df.select_dtypes(include=['object']).columns
                for col in cols:
                    df[col] = df[col].apply(lambda x: translate_block(x, st.session_state.src_code, st.session_state.target_code))
                out = io.BytesIO()
                if file.name.endswith(".csv"):
                    df.to_csv(out, index=False); mime = "text/csv"
                else:
                    df.to_excel(out, index=False); mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                data = out.getvalue()
            
            results[file.name] = (data, mime)
            st.session_state.history.append({"name": file.name, "data": data, "mime": mime, "target_name": st.session_state.target_name, "timestamp": time.time()})
            status.update(label=f"✅ {file.name} Finished", state="complete")
            progress_bar.progress((idx + 1) / len(st.session_state.files))

    st.session_state.final_results = results
    st.session_state.step = "results"; st.rerun()

elif st.session_state.step == "results":
    st.balloons()
    st.title("🎉 Translation Ready")
    for name, (data, mime) in st.session_state.final_results.items():
        trigger_auto_download(f"translet_{name}", data, mime)
    
    grid = st.columns(3)
    for i, (name, (data, mime)) in enumerate(st.session_state.final_results.items()):
        with grid[i % 3]:
            st.info(f"📄 {name}")
            st.download_button("Download Again", data, f"translet_{name}", mime, key=f"res_{i}")
    
    st.divider()
    if st.button("✨ Upload New Translation"):
        st.session_state.step = "upload"; st.rerun()